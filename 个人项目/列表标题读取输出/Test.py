# 研发者：时间遗忘
# 开发时间：2023/10/20 8:50
#遍历所有的四级标题，并打印输出到word
import docx

# 打开要读取的Word文档
doc = docx.Document('o2o项目需求文档.docx')  # 将'your_document.docx'替换为您的文档路径

# 创建一个新的Word文档，用于存放四级标题
output_doc = docx.Document()

# 遍历文档内容，提取四级标题
fourth_level_titles = []
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith('Heading 4'):
        fourth_level_titles.append(paragraph.text)

# 创建一个新的表格
table = output_doc.add_table(rows=1, cols=1)

# 向表格中添加标题
table.cell(0, 0).text = "Fourth Level Titles"

# 向表格中逐行添加四级标题
for title in fourth_level_titles:
    row = table.add_row().cells
    row[0].text = title

# 保存输出的Word文档
output_doc.save('output_table.docx')  # 将'output_table.docx'替换为您希望保存的文件名

print("四级标题已提取并保存到output_table.docx文件中。")

