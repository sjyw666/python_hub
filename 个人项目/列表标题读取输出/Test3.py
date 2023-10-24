# 研发者：时间遗忘
# 开发时间：2023/10/20 8:59
#在二的基础上添加了三级标题在前面一列
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# 打开要读取的Word文档
doc = Document('o2o项目需求文档.docx')  # 将'your_document.docx'替换为您的文档路径

# 创建一个新的Word文档
output_doc = Document()

# 创建一个两列的表格，用于存放三级标题和四级标题
table = output_doc.add_table(rows=1, cols=2)

# 设置表格样式
table.style = 'Table Grid'

# 设置表格的列宽度，根据需要进行调整
table.autofit = False
table.columns[0].width = Pt(200)  # 第一列的列宽
table.columns[1].width = Pt(200)  # 第二列的列宽

# 向表格中添加标题行
table.cell(0, 0).text = "三级标题"
table.cell(0, 1).text = "四级标题"

# 初始化三级标题和四级标题
third_level_title = ""
fourth_level_title = ""

# 读取文档中的标题并将其写入表格
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith('Heading 3'):
        third_level_title = paragraph.text
    elif paragraph.style.name.startswith('Heading 4'):
        fourth_level_title = paragraph.text
        # 向表格中添加一行，将三级标题放在第一列，四级标题放在第二列
        row = table.add_row().cells
        row[0].text = third_level_title
        row[1].text = fourth_level_title

# 保存输出的Word文档
output_doc.save('output_table3.docx')  # 将'output_table.docx'替换为您希望保存的文件名

print("三级标题和四级标题已提取并保存到output_table.docx文件中。")
