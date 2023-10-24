# 研发者：时间遗忘
# 开发时间：2023/10/20 8:54
#在1的基础上增加了表格
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# 打开要读取的Word文档
doc = Document('o2o项目需求文档.docx')  # 将'your_document.docx'替换为您的文档路径

# 创建一个新的Word文档
output_doc = Document()

# 创建一个表格，每行一个四级标题
table = output_doc.add_table(rows=1, cols=1)

# 设置表格样式
table.style = 'Table Grid'

# 设置表格的列宽度，根据需要进行调整
table.autofit = False
table.columns[0].width = Pt(400)  # 调整列宽度

# 向表格中添加标题行
table.cell(0, 0).text = "四级标题"

# 读取文档中的四级标题并将其写入表格
for paragraph in doc.paragraphs:
    if paragraph.style.name.startswith('Heading 4'):
        row = table.add_row().cells
        row[0].text = paragraph.text

# 保存输出的Word文档
output_doc.save('output_table2.docx')  # 将'output_table.docx'替换为您希望保存的文件名

print("四级标题已提取并保存到output_table.docx文件中。")
